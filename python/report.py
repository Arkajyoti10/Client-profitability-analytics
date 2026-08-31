import os
from pathlib import Path
from datetime import date

from dotenv import load_dotenv
import pandas as pd
import numpy as np
from groq import Groq
from docx import Document

from extract import get_client_profitability
from segment import segment_clients
from flag import flag_underperforming_clients

# Setup (runs once on import)
BASE_DIR = Path(__file__).resolve().parent.parent
load_dotenv(BASE_DIR / ".env")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise EnvironmentError("Missing GROQ_API_KEY. Check your .env file.")

client = Groq(api_key=GROQ_API_KEY)

SYSTEM_PROMPT = """You are a data analyst at a professional services firm writing an internal memo for a managing partner.

You will be given a table of clients flagged as underperforming on profitability, split into two groups:
- "Discount-driven": fee is below standard value of work performed (realization rate under 100%)
- "Cost-driven": fee is at/above standard value, but staffing cost is eroding margin

Respond with exactly two sections, each on its own line starting with these exact labels:

SUMMARY: <2-3 sentence overview of the situation across both groups, including the split between discount-driven and cost-driven counts>
RECOMMENDATION: <2-3 sentence closing recommendation on which group to prioritize first and why>

Do not include any other text, headers, or the data table itself — those are handled separately. Be concise and professional.
"""

# Root-cause classification
def classify_root_cause(flagged_df: pd.DataFrame) -> pd.DataFrame:
    flagged_df["root_cause"] = np.where(
        flagged_df["realization_rate"] < 100,
        "Discount-driven",
        "Cost-driven"
    )
    return flagged_df

# LLM prompt construction + call
def build_user_message(classified_df: pd.DataFrame) -> str:
    table_string = classified_df[
        ["client_id", "net_margin_%", "realization_rate", "root_cause"]
    ].to_markdown(index=False)

    return f"""Here is the flagged client data:

{table_string}

Write the SUMMARY and RECOMMENDATION sections as instructed."""


def parse_llm_sections(response_text: str) -> dict:
    summary = ""
    recommendation = ""
    if "RECOMMENDATION:" in response_text:
        summary_part, recommendation_part=response_text.split("RECOMMENDATION:",1)
        summary=summary_part.replace("SUMMARY:","").strip()
        recommendation=recommendation_part.strip()
    else:
        summary=response_text.replace("SUMMARY:","").strip()
    return {"summary": summary, "recommendation": recommendation}


def generate_report(classified_df: pd.DataFrame) -> dict:
    user_message = build_user_message(classified_df)

    try:
        response = client.chat.completions.create(
            model="openai/gpt-oss-120b",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_message}
            ]
        )
    except Exception as e:
        raise RuntimeError(f"Failed to generate report: {e}")

    return parse_llm_sections(response.choices[0].message.content)

# Word document building
def _add_client_table(doc: Document, df: pd.DataFrame):
    if df.empty:
        doc.add_paragraph("None.")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Client ID"
    hdr_cells[1].text = "Net Margin %"
    hdr_cells[2].text = "Realization Rate %"

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["client_id"])
        cells[1].text = f"{row['net_margin_%']:.2f}"
        cells[2].text = f"{row['realization_rate']:.2f}"

#generating the document report
def build_docx_report(classified_df: pd.DataFrame, sections: dict, output_path: Path):
    doc = Document()

    doc.add_heading("Under-performing Client Portfolio – Root-Cause Summary & Action Plan", level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().strftime('%d %b %Y')}\n")
    meta.add_run("Prepared by: Client Profitability Analytics Pipeline")

    doc.add_heading("1. Overall Summary", level=2)
    doc.add_paragraph(sections["summary"])

    discount_df = classified_df[classified_df["root_cause"] == "Discount-driven"]
    cost_df = classified_df[classified_df["root_cause"] == "Cost-driven"]

    doc.add_heading("2. Discount-Driven Clients", level=2)
    _add_client_table(doc, discount_df)

    doc.add_heading("3. Cost-Driven Clients", level=2)
    _add_client_table(doc, cost_df)

    doc.add_heading("4. Recommendation", level=2)
    doc.add_paragraph(sections["recommendation"])

    doc.save(output_path)

#run the pipeline
if __name__ == "__main__":
    df = get_client_profitability()
    segmented = segment_clients(df)
    flagged = flag_underperforming_clients(segmented)
    classified = classify_root_cause(flagged)

    sections = generate_report(classified)

    output_path = BASE_DIR / "deliverables" / "profitability_report.docx"
    build_docx_report(classified, sections, output_path)

    print(f"Report saved to {output_path}")